"""
文档解析技能处理器
功能：PDF/Word/Excel/CSV/图片解析
"""
import json
from pathlib import Path


class DocParserSkill:
    """文档解析技能"""

    @staticmethod
    def parse_pdf(filepath: str) -> dict:
        """解析PDF文件"""
        try:
            from PyPDF2 import PdfReader
            path = Path(filepath)
            if not path.exists():
                return {"error": f"文件不存在: {filepath}"}

            reader = PdfReader(str(path))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)

            return {
                "success": True,
                "filename": path.name,
                "total_pages": len(reader.pages),
                "content": "\n".join(pages),
                "pages": pages
            }
        except Exception as e:
            return {"error": f"PDF解析失败: {str(e)}"}

    @staticmethod
    def parse_docx(filepath: str) -> dict:
        """解析Word文档"""
        try:
            from docx import Document
            path = Path(filepath)
            if not path.exists():
                return {"error": f"文件不存在: {filepath}"}

            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # 提取表格
            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                tables.append(rows)

            return {
                "success": True,
                "filename": path.name,
                "paragraphs": paragraphs,
                "tables": tables,
                "content": "\n".join(paragraphs)
            }
        except Exception as e:
            return {"error": f"Word解析失败: {str(e)}"}

    @staticmethod
    def parse_xlsx(filepath: str, sheet_name: str = None) -> dict:
        """解析Excel文件"""
        try:
            import pandas as pd
            path = Path(filepath)
            if not path.exists():
                return {"error": f"文件不存在: {filepath}"}

            if sheet_name:
                df = pd.read_excel(str(path), sheet_name=sheet_name)
            else:
                # 读取所有sheet
                xls = pd.ExcelFile(str(path))
                sheets = {}
                for sheet in xls.sheet_names:
                    df = pd.read_excel(str(path), sheet_name=sheet)
                    sheets[sheet] = {
                        "columns": df.columns.tolist(),
                        "rows": len(df),
                        "data": df.head(1000).to_dict(orient="records")
                    }
                return {
                    "success": True,
                    "filename": path.name,
                    "sheets": sheets
                }

            return {
                "success": True,
                "filename": path.name,
                "sheet": sheet_name,
                "columns": df.columns.tolist(),
                "rows": len(df),
                "data": df.head(1000).to_dict(orient="records")
            }
        except Exception as e:
            return {"error": f"Excel解析失败: {str(e)}"}

    @staticmethod
    def parse_csv(filepath: str) -> dict:
        """解析CSV文件"""
        try:
            import pandas as pd
            path = Path(filepath)
            if not path.exists():
                return {"error": f"文件不存在: {filepath}"}

            df = pd.read_csv(str(path))
            return {
                "success": True,
                "filename": path.name,
                "columns": df.columns.tolist(),
                "rows": len(df),
                "data": df.head(1000).to_dict(orient="records")
            }
        except Exception as e:
            return {"error": f"CSV解析失败: {str(e)}"}

    @staticmethod
    def parse_text(filepath: str) -> dict:
        """解析纯文本文件"""
        path = Path(filepath)
        if not path.exists():
            return {"error": f"文件不存在: {filepath}"}

        try:
            content = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            try:
                content = path.read_text(encoding="gbk")
            except Exception as e:
                return {"error": f"文本编码识别失败: {str(e)}"}

        return {
            "success": True,
            "filename": path.name,
            "content": content,
            "lines": len(content.split("\n")),
            "size": len(content)
        }


# 技能入口函数
def execute(action: str, params: dict) -> dict:
    """技能执行入口"""
    skill = DocParserSkill()

    actions = {
        "pdf": skill.parse_pdf,
        "docx": skill.parse_docx,
        "xlsx": skill.parse_xlsx,
        "csv": skill.parse_csv,
        "txt": skill.parse_text,
    }

    if action not in actions:
        return {"error": f"不支持的格式: {action}"}

    return actions[action](params.get("filepath", ""))